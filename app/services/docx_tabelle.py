"""Helper riusabili per tabelle Word (stile modulo ufficiale)."""
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table, _Cell, _Row


def tabella_griglia(doc: DocumentObject, rows: int, cols: int, col_cm: float = 9.29) -> Table:
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    t.autofit = True
    for row in t.rows:
        for cell in row.cells:
            cell.width = Cm(col_cm)
    return t


def cella_multiline(cell: _Cell, linee: list[str], font_pt: float = 14, bold: bool = False) -> None:
    """Sostituisce il contenuto della cella con testo a più righe (break Word)."""
    p = cell.paragraphs[0]
    p.clear()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    for i, linea in enumerate(linee):
        if i:
            p.add_run().add_break()
        run = p.add_run(linea)
        run.font.size = Pt(font_pt)
        run.bold = bold


def riga_altezza_min_cm(row: _Row, altezza_cm: float, *, esatta: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for child in list(tr_pr):
        if child.tag == qn("w:trHeight"):
            tr_pr.remove(child)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(altezza_cm * 567)))
    h.set(qn("w:hRule"), "exact" if esatta else "atLeast")
    tr_pr.append(h)


def cella_margini_pt(
    cell: _Cell,
    *,
    top: float = 8,
    bottom: float = 8,
    left: float = 8,
    right: float = 8,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:tcMar"):
            tc_pr.remove(child)
    tc_mar = OxmlElement("w:tcMar")
    for nome, val_pt in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{nome}")
        node.set(qn("w:w"), str(int(val_pt * 20)))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)