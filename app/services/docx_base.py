"""Elementi comuni Word: intestazione ente (logo, comune, provincia)."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.extensions import db
from app.models.ente import EnteSettings
from app.services.logo_ente import logo_ente_path


def _riga_comune(ente: EnteSettings | None) -> str:
    if not ente:
        return ""
    denom = (ente.denominazione or "").strip()
    if denom:
        return denom.upper()
    comune = (ente.comune or "").strip()
    if comune:
        return f"COMUNE DI {comune}".upper()
    return ""


def _riga_provincia(ente: EnteSettings | None) -> str:
    if not ente:
        return ""
    prov = (ente.provincia or "").strip()
    if not prov:
        return ""
    return f"(PROVINCIA DI {prov})".upper()


def applica_intestazione_ente(
    doc: Document,
    *,
    logo_cm: float = 3.0,
    font_comune: float = 12,
    font_provincia: float = 11,
) -> None:
    """Header come nel modulo ufficiale: logo centrato + comune + provincia."""
    ente = db.session.get(EnteSettings, 1)
    logo = logo_ente_path()
    comune = _riga_comune(ente)
    provincia = _riga_provincia(ente)

    for section in doc.sections:
        header = section.header
        p0 = header.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        p0.clear()
        if logo is not None:
            run = p0.add_run()
            run.add_picture(str(logo), width=Cm(logo_cm))

        if comune:
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(comune)
            r.bold = True
            r.font.size = Pt(font_comune)

        if provincia:
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(provincia)
            r.bold = True
            r.font.size = Pt(font_provincia)
