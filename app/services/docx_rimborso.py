"""Genera il modulo Word di richiesta rimborso da un buono economale."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.config import INSTANCE_DIR
from app.models.buono import BuonoEconomale
from app.models.economo import EconomoSettings
from app.services.docx_base import applica_intestazione_ente
from app.services.docx_tabelle import (
    cella_margini_pt,
    cella_multiline,
    riga_altezza_min_cm,
    tabella_griglia,
)
from app.services.economo_testo import nome_da_economo
from app.services.firme_rimborso import LINEA_FIRMA, linee_firma_modulo
from app.services.numero_display import formato_numero_sezionale

_FONT = 11
_FONT_TITOLO = 13


def _eur(v) -> str:
    try:
        return f"{float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return "—"


def _data_it(d) -> str:
    if not d:
        return "____ / ____ / ______"
    return d.strftime("%d / %m / %Y")


def _applica_margini(tab, **kwargs) -> None:
    for row in tab.rows:
        for cell in row.cells:
            cella_margini_pt(cell, **kwargs)


def _parti_blocco(blocco: list[str]) -> tuple[str, str, str]:
    etichetta = blocco[0] if blocco else ""
    linea = blocco[-1] if blocco else LINEA_FIRMA
    nome = blocco[1] if len(blocco) > 2 else ""
    return etichetta, nome, linea


def _para(
    doc: Document,
    text: str,
    *,
    center: bool = False,
    bold: bool = False,
    size: float | None = _FONT,
    space_before: float = 0,
    space_after: float = 4,
):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    return p


def genera_docx_rimborso(b: BuonoEconomale) -> Path:
    out_dir = INSTANCE_DIR / "buoni_docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"modulo_rimborso_{b.anno}_{b.numero_progressivo:04d}.docx"

    num = formato_numero_sezionale(b)
    importo = _eur(b.importo_speso or b.importo_autorizzato)
    data = _data_it(b.data_buono)
    richiedente = (b.richiedente or "").strip() or "___________________________________________________"
    ufficio = (b.ufficio_richiedente or "").strip() or "____________________________"
    responsabile = (getattr(b, "responsabile_ufficio", None) or "").strip()
    causale = (b.causale or "").strip() or ("_" * 80)

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        # Il logo sta nell'header: il margine alto deve coprirlo, sennò Word va a 2 pagine.
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.header_distance = Cm(0.3)
        section.footer_distance = Cm(0.4)

    applica_intestazione_ente(doc, logo_cm=1.6, font_comune=11, font_provincia=10)

    _para(
        doc,
        "RICHIESTA DI RIMBORSO SPESA ECONOMALE",
        center=True,
        bold=True,
        size=_FONT_TITOLO,
        space_after=8,
    )

    tab_num = tabella_griglia(doc, 1, 2)
    cella_multiline(tab_num.rows[0].cells[0], [f"Buono economale n. {num}"], font_pt=_FONT)
    cella_multiline(tab_num.rows[0].cells[1], [f"Data {data}"], font_pt=_FONT)
    riga_altezza_min_cm(tab_num.rows[0], 1.0, esatta=True)
    _applica_margini(tab_num, top=6, bottom=6, left=8, right=8)

    _para(
        doc,
        f"Il/La sottoscritto/a {richiedente}, dell’Ufficio {ufficio}, "
        f"chiede il rimborso di € {importo} per:",
        space_before=8,
        space_after=6,
    )

    tab_cau = tabella_griglia(doc, 2, 1, col_cm=18.6)
    cella_multiline(
        tab_cau.rows[0].cells[0],
        ["Causale / oggetto della spesa"],
        font_pt=_FONT,
        bold=True,
    )
    cella_multiline(tab_cau.rows[1].cells[0], [causale], font_pt=_FONT)
    riga_altezza_min_cm(tab_cau.rows[0], 0.8, esatta=True)
    riga_altezza_min_cm(tab_cau.rows[1], 2.8, esatta=True)
    _applica_margini(tab_cau, top=6, bottom=6, left=8, right=8)

    eco = EconomoSettings.query.get(1)
    blocchi = linee_firma_modulo(
        b.richiedente,
        b.ufficio_richiedente,
        responsabile,
        nome_da_economo(eco),
    )
    _para(doc, f"Data {data}", space_before=10, space_after=8)
    n = len(blocchi)
    col_cm = 6.2 if n == 3 else 9.29
    tab_firme = tabella_griglia(doc, 2, n, col_cm=col_cm)
    for i, blocco in enumerate(blocchi):
        etichetta, nome, linea = _parti_blocco(blocco)
        testa = [etichetta, nome] if nome else [etichetta]
        cella_multiline(tab_firme.rows[0].cells[i], testa, font_pt=_FONT, bold=True)
        cella_multiline(tab_firme.rows[1].cells[i], [linea], font_pt=_FONT)
        tab_firme.rows[1].cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    riga_altezza_min_cm(tab_firme.rows[0], 1.5, esatta=True)
    riga_altezza_min_cm(tab_firme.rows[1], 4.2, esatta=True)
    _applica_margini(tab_firme, top=6, bottom=10, left=6, right=6)

    doc.save(str(path))
    return path
